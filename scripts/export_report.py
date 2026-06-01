from __future__ import annotations

import argparse
import html
import re
import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_SIZES = {
    "letter": (8.5, 11.0),
    "a4": (8.27, 11.69),
}

DIAGRAM_CAPTIONS = [
    "Figure 3.1: End-to-End Smart Eye Monitoring Pipeline",
    "Figure 3.2: Smart Eye Layered System Architecture",
    "Figure 3.3: Object Detection Training and ONNX Deployment Workflow",
    "Figure 3.4: Feature Representation and Rule Attribute Generation",
    "Figure 3.5: Runtime Inference Sequence for Camera and Video Frames",
    "Figure 3.6: Face Recognition and Liveness Verification Workflow",
    "Figure 3.7: Rule Evaluation and Alarm Action Flow",
    "Figure 3.8: Alarm Escalation State Flow",
    "Figure 3.9: Smart Eye Database Entity Relationship Overview",
    "Figure 3.10: Analytics and PDF Reporting Workflow",
    "Figure 3.11: User Interface Navigation Structure",
    "Figure 4.1: Dataset Preparation and Training Data Flow",
]

CODE_LABELS = {
    "python": "Implementation excerpt",
    "bash": "Build command",
    "powershell": "PowerShell command",
    "text": "Rule condition example",
}


CSS = """
@page {
  size: Letter;
  margin: 1in;
}
body {
  font-family: "Times New Roman", serif;
  font-size: 12pt;
  line-height: 1.45;
  color: #111;
}
h1 {
  font-size: 22pt;
  text-align: center;
  margin: 0 0 14pt;
}
h2 {
  font-size: 18pt;
  margin: 18pt 0 12pt;
  break-after: avoid;
  page-break-after: avoid;
}
h3 {
  font-size: 14pt;
  margin: 18pt 0 8pt;
  break-after: avoid;
  page-break-after: avoid;
}
h4 {
  font-size: 12.5pt;
  margin: 14pt 0 6pt;
  break-after: avoid;
  page-break-after: avoid;
}
p {
  margin: 0 0 9pt;
  text-align: justify;
  orphans: 3;
  widows: 3;
}
ul, ol {
  margin-top: 0;
  margin-bottom: 9pt;
}
li {
  margin-bottom: 4pt;
}
table {
  width: 100%;
  border-collapse: collapse;
  margin: 8pt 0 12pt;
  break-inside: avoid;
  page-break-inside: avoid;
  font-size: 10pt;
}
th, td {
  border: 0.75pt solid #444;
  padding: 5pt 6pt;
  vertical-align: top;
}
th {
  background: #eceff3;
  font-weight: bold;
}
pre {
  position: relative;
  font-family: "Courier New", monospace;
  font-size: 8.2pt;
  line-height: 1.35;
  white-space: pre-wrap;
  background: #f8fafc;
  border: 0.75pt solid #d6dbe2;
  border-left: 2.5pt solid #6b7280;
  border-radius: 3pt;
  padding: 8pt;
  margin: 3pt 0 12pt;
  break-inside: avoid;
  page-break-inside: avoid;
  color: #1d2733;
}
.code-block {
  break-inside: avoid;
  page-break-inside: avoid;
}
code {
  font-family: "Courier New", monospace;
  font-size: 9pt;
  color: #1d2733;
}
.code-label {
  display: block;
  margin: 8pt 0 3pt;
  padding: 0;
  color: #333;
  font-family: "Times New Roman", serif;
  font-size: 10pt;
  font-style: italic;
  font-weight: bold;
}
.code-keyword {
  color: #1f4e79;
  font-weight: bold;
}
.code-string {
  color: #2f6f4e;
}
.code-comment {
  color: #667085;
  font-style: italic;
}
.placeholder {
  color: #b42318;
  font-weight: bold;
  font-style: italic;
}
blockquote {
  margin: 8pt 0 12pt 18pt;
  padding-left: 10pt;
  border-left: 3pt solid #c8c8c8;
  color: #333;
}
hr {
  border: none;
  height: 0;
  margin: 0;
  break-after: page;
  page-break-after: always;
}
figure.diagram {
  margin: 10pt 0 14pt;
  padding: 8pt;
  border: 0.75pt solid #d6dbe2;
  background: #fbfcfe;
  break-inside: avoid;
  page-break-inside: avoid;
}
.mermaid {
  text-align: center;
  overflow: visible;
}
.mermaid svg {
  max-width: 100%;
  height: auto !important;
}
figcaption {
  margin-top: 6pt;
  font-size: 10pt;
  color: #444;
  text-align: center;
  font-style: italic;
}
.math {
  text-align: center;
  margin: 10pt 0 14pt;
  font-size: 13pt;
  break-inside: avoid;
  page-break-inside: avoid;
}
@media print {
  a {
    color: #111;
    text-decoration: none;
  }
}
"""


MERMAID_SCRIPT = """
<script type="module">
import mermaid from "https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs";
window.__mermaidReady = false;
mermaid.initialize({
  startOnLoad: false,
  securityLevel: "loose",
  theme: "default",
  flowchart: { useMaxWidth: true, htmlLabels: true },
  sequence: { useMaxWidth: true }
});
try {
  await mermaid.run({ querySelector: ".mermaid" });
  window.__mermaidReady = true;
} catch (error) {
  console.error("Mermaid render failed", error);
  window.__mermaidReady = "error";
}
</script>
"""

MATHJAX_SCRIPT = """
<script>
window.__mathReady = false;
window.MathJax = {
  tex: {
    inlineMath: [["\\\\(", "\\\\)"]],
    displayMath: [["\\\\[", "\\\\]"]]
  },
  svg: { fontCache: "global" },
  startup: {
    pageReady: () => MathJax.startup.defaultPageReady().then(() => {
      window.__mathReady = true;
    })
  }
};
</script>
<script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-svg.js"></script>
"""


def replace_math_blocks(markdown_text: str) -> str:
    index = 0

    def repl(match: re.Match[str]) -> str:
        nonlocal index
        equation = match.group(1).strip()
        html_block = f'\n<div class="math display" data-equation-index="{index}">\\[{equation}\\]</div>\n'
        index += 1
        return html_block

    return re.sub(r"\$\$\s*(.*?)\s*\$\$", repl, markdown_text, flags=re.DOTALL)


def highlight_code(code_text: str, language: str) -> str:
    escaped = html.escape(code_text)

    placeholders: list[str] = []

    def protect(value: str) -> str:
        token = f"@@CODETOKEN{len(placeholders)}@@"
        placeholders.append(value)
        return token

    if language in {"python", "bash", "powershell"}:
        escaped = re.sub(
            r"(?m)(#.*)$",
            lambda match: protect(f'<span class="code-comment">{match.group(1)}</span>'),
            escaped,
        )
    escaped = re.sub(
        r"(&quot;.*?&quot;|'.*?')",
        lambda match: protect(f'<span class="code-string">{match.group(1)}</span>'),
        escaped,
    )
    if language == "python":
        escaped = re.sub(
            r"\b(def|class|return|if|elif|else|for|while|try|except|with|as|import|from|None|True|False)\b",
            r'<span class="code-keyword">\1</span>',
            escaped,
        )
    for idx, value in enumerate(placeholders):
        escaped = escaped.replace(f"@@CODETOKEN{idx}@@", value)
    return escaped


def infer_code_listing_title(code_text: str, language: str, fallback_index: int) -> str:
    compact = re.sub(r"\s+", " ", code_text)
    checks = [
        ("export_model.export", "Listing 3.1: ONNX Export Configuration"),
        ("train_args = dict", "Listing 5.1: YOLO11s Training Configuration"),
        ("model.export(", "Listing 5.2: ONNX Export Command"),
        ("ONNXObjectModel", "Listing 5.3: ONNX Model Loading"),
        ("DATA_DIR = BASE_DIR", "Listing 5.4: Application Startup Initialization"),
        ("detector.process_frame", "Listing 5.5: Camera Frame Inference"),
        ("pipeline.handle_result", "Listing 5.6: Pipeline Result Handling"),
        ("NO-Hardhat = true", "Listing 5.7: Example Rule Conditions"),
        ("def start_camera", "Listing 5.8: Camera Thread Startup"),
        ("def evaluate_rules", "Listing 5.9: Rule Evaluation Loop"),
        ('"dashboard": PageSpec', "Listing 5.10: Dashboard Page Registration"),
        ("db.add_detection_log", "Listing 5.11: Detection Log Persistence"),
        ("def _hash_password", "Listing 5.12: PBKDF2 Password Hashing"),
        ('atype == "email"', "Listing 5.13: Alarm Action Dispatch"),
        ("def generate_report", "Listing 5.15: Analytics Report Generation"),
        ("stats_engine.get_summary", "Listing 5.14: Analytics Query Generation"),
        ("python -m nuitka", "Listing 5.16: Windows Build Command"),
    ]
    for needle, title in checks:
        if needle in compact:
            return title
    fallback = CODE_LABELS.get(language, "Implementation excerpt")
    return f"Listing {fallback_index + 1}: {fallback}"


def replace_mermaid_blocks(body: str) -> tuple[str, int]:
    soup = BeautifulSoup(body, "html.parser")
    count = 0
    for code in list(soup.find_all("code")):
        classes = code.get("class", [])
        if "language-mermaid" not in classes:
            continue
        pre = code.find_parent("pre")
        if pre is None:
            continue
        figure = soup.new_tag("figure")
        figure["class"] = "diagram"
        diagram = soup.new_tag("div")
        diagram["class"] = "mermaid"
        diagram["data-diagram-index"] = str(count)
        diagram.string = code.get_text()
        figure.append(diagram)
        caption = soup.new_tag("figcaption")
        caption.string = DIAGRAM_CAPTIONS[count] if count < len(DIAGRAM_CAPTIONS) else f"Figure: Smart Eye workflow diagram {count + 1}"
        figure.append(caption)
        pre.replace_with(figure)
        count += 1
    return str(soup), count


def style_code_blocks(body: str) -> str:
    soup = BeautifulSoup(body, "html.parser")
    count = 0
    for code in list(soup.find_all("code")):
        pre = code.find_parent("pre")
        if pre is None:
            continue
        classes = code.get("class", [])
        language = ""
        for cls in classes:
            if cls.startswith("language-"):
                language = cls.replace("language-", "", 1)
                break
        if language == "mermaid":
            continue
        original_code = code.get_text()
        label = infer_code_listing_title(original_code, language, count)
        wrapper = soup.new_tag("div")
        wrapper["class"] = "code-block"
        wrapper["data-code-index"] = str(count)
        wrapper["data-code-title"] = label
        label_tag = soup.new_tag("div")
        label_tag["class"] = "code-label"
        label_tag.string = label
        code["class"] = classes + [f"code-language-{language or 'plain'}"]
        code.clear()
        highlighted = BeautifulSoup(highlight_code(original_code, language), "html.parser")
        for child in list(highlighted.contents):
            code.append(child)
        pre.replace_with(wrapper)
        wrapper.append(label_tag)
        wrapper.append(pre)
        count += 1
    return str(soup)


def markdown_to_html(markdown_text: str) -> str:
    markdown_text = replace_math_blocks(markdown_text)
    body = markdown.markdown(
        markdown_text,
        extensions=[
            "extra",
            "fenced_code",
            "sane_lists",
            "tables",
            "toc",
        ],
        output_format="html5",
    )
    body = style_code_blocks(body)
    body, mermaid_count = replace_mermaid_blocks(body)
    mermaid_script = MERMAID_SCRIPT if mermaid_count else ""
    math_script = MATHJAX_SCRIPT if 'class="math display"' in body else ""
    return f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>Smart Eye Final Report</title>
<style>{CSS}</style>
</head>
<body>
{body}
{mermaid_script}
{math_script}
</body>
</html>
"""


def write_html(markdown_path: Path, output_path: Path) -> str:
    text = markdown_path.read_text(encoding="utf-8")
    html_text = markdown_to_html(text)
    output_path.write_text(html_text, encoding="utf-8")
    return html_text


def normalize_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip()


def inline_text(node: Tag) -> str:
    text = node.get_text(" ", strip=True)
    return normalize_text(text)


def configure_docx(document: Document, page_size: str) -> None:
    width, height = PAGE_SIZES[page_size]
    section = document.sections[0]
    section.page_width = Inches(width)
    section.page_height = Inches(height)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for name, size in (
        ("Title", 22),
        ("Heading 1", 18),
        ("Heading 2", 14),
        ("Heading 3", 12.5),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)


def add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    text = normalize_text(text)
    if not text:
        return
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run(text)


def add_word_field(paragraph, field_code: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)

    if placeholder:
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        run._r.append(separate)
        result = paragraph.add_run(placeholder)
        result.italic = True
        result.font.color.rgb = RGBColor(0x66, 0x70, 0x85)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_tc_field(paragraph, text: str, identifier: str) -> None:
    safe_text = normalize_text(text).replace('"', "'")
    if not safe_text:
        return
    add_word_field(paragraph, f'TC "{safe_text}" \\f {identifier} \\l 1')


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_auto_field_block(document: Document, field_code: str, placeholder: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    add_word_field(paragraph, field_code, placeholder=placeholder)


def add_inline_runs(
    paragraph,
    node: Tag,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
    placeholder: bool = False,
) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child).replace("\xa0", " ")
            if not text:
                continue
            run = paragraph.add_run(text)
            run.bold = bold or placeholder
            run.italic = italic or placeholder
            if placeholder:
                run.font.color.rgb = RGBColor(0xB4, 0x23, 0x18)
            elif code:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1D, 0x27, 0x33)
            elif bold and italic:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            continue

        if not isinstance(child, Tag):
            continue
        if child.name == "br":
            paragraph.add_run().add_break()
            continue
        classes = child.get("class", []) if isinstance(child, Tag) else []
        add_inline_runs(
            paragraph,
            child,
            bold=bold or child.name in {"strong", "b"},
            italic=italic or child.name in {"em", "i"},
            code=code or child.name == "code",
            placeholder=placeholder or "placeholder" in classes,
        )


def add_rich_paragraph(document: Document, node: Tag, style: str | None = None) -> None:
    if not inline_text(node):
        return
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(8)
    add_inline_runs(paragraph, node)
    text = inline_text(node)
    if re.match(r"^Figure\s+(?:[A-Z]\.)?\d", text):
        add_tc_field(paragraph, text, "F")
    elif re.match(r"^Table\s+(?:[A-Z]\.)?\d", text):
        add_tc_field(paragraph, text, "T")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 120, start: int = 160, bottom: int = 120, end: int = 160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_code_block(document: Document, text: str, language: str = "", label: str | None = None) -> None:
    if label is None:
        label = "Diagram source fallback" if language == "mermaid" else infer_code_listing_title(text, language, 0)

    if label:
        label_paragraph = document.add_paragraph()
        label_paragraph.paragraph_format.keep_with_next = True
        label_paragraph.paragraph_format.space_before = Pt(6)
        label_paragraph.paragraph_format.space_after = Pt(2)
        label_run = label_paragraph.add_run(label)
        label_run.bold = True
        label_run.italic = True
        label_run.font.name = "Times New Roman"
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = True
    row = table.rows[0]
    keep_table_row_together(row)
    cell = row.cells[0]
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    lines = text.rstrip("\n").splitlines() or [""]
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1D, 0x27, 0x33)
        if idx < len(lines) - 1:
            run.add_break()

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


TABLE_CAPTION_BY_HEADERS = {
    ("Monitoring Need", "Traditional Camera System", "Smart Eye Approach"): "Table 1.1: Traditional Surveillance Compared with Smart Eye",
    ("Area", "Included in Smart Eye", "Not Included in Current Scope"): "Table 1.2: Project Scope Boundaries",
    ("Objective", "Related Deliverable"): "Table 1.3: Research Objectives and Deliverables",
    ("Output Type", "Examples"): "Table 1.4: Expected System Outputs",
    ("Research Area", "Contribution to Smart Eye"): "Table 2.1: Literature Areas and Project Relevance",
    ("Requirement", "Importance for Smart Eye"): "Table 2.2: Object Detection Requirements",
    ("Detected Condition", "Possible Rule", "Possible System Response"): "Table 2.3: PPE Detection and System Response Mapping",
    ("Stage", "Purpose"): "Table 2.4: Face Recognition Pipeline Stages",
    ("Question", "Smart Eye Component"): "Table 2.5: Detection-to-Action Integration Questions",
    ("Layer", "Contribution"): "Table 2.6: Project Contribution Layers",
    ("Study / Source", "Technique", "Dataset / Domain", "Main Result", "Limitation", "Relation to Smart Eye"): "Table 2.7: Comparison of Previous Studies",
    ("Technology", "Role in Smart Eye", "Strength", "Limitation"): "Table 2.8: Technology Comparison",
    ("Requirement Type", "Minimum Specification", "Recommended Specification"): "Table 3.1: Hardware Requirements",
    ("Category", "Technology Used", "Purpose"): "Table 3.2: Software Requirements",
    ("Requirement Concern", "Design Response in Smart Eye"): "Table 3.3: Requirement Concerns and Design Responses",
    ("Methodology Stage", "Input", "Processing", "Output"): "Table 3.4: Proposed Methodology Stages",
    ("Architectural View", "Main Concern", "Example in Smart Eye"): "Table 3.5: Architectural Views",
    ("User-Facing Module", "Code Module", "Main Responsibility"): "Table 3.6: Main Application Modules",
    ("Training Choice", "Reason"): "Table 3.7: Object Detection Training Choices",
    ("Level", "Description", "Example Output"): "Table 3.8: Runtime Detection Pipeline Levels",
    ("Configuration Element", "Effect"): "Table 3.9: Face Recognition Configuration Effects",
    ("Rule Aspect", "Purpose"): "Table 3.10: Rule Configuration Elements",
    ("Workflow", "Database Role"): "Table 3.11: Database Workflow Support",
    ("Analytics Output", "Data Used", "Purpose"): "Table 3.12: Analytics Outputs",
    ("Data Category", "Created By", "Used For", "Stored In"): "Table 4.1: Dataset and Operational Data Categories",
    ("Field", "Description"): "Table 4.2: Object Detection Dataset Source",
    ("Class Name", "Type", "Description", "Expected Value / Range"): "Table 4.3: Dataset Classes and Features",
    ("Class", "Example Application Use"): "Table 4.4: Dataset Classes and Application Use",
    ("Augmentation", "Expected Benefit"): "Table 4.5: Data Augmentation Benefits",
    ("Class", "Before Balancing Boxes", "After Balancing Boxes", "Added Training Copies", "Class Weight"): "Table 4.6: Class Balancing and Training Weights",
    ("Split", "Image Count", "Percentage", "Purpose"): "Table 4.6: Dataset Split Plan",
    ("Dataset Item", "Image Count", "Percentage / Notes", "Purpose"): "Table 4.7: Dataset Processing and Split Counts",
    ("Data Type", "Storage Location / Table", "Purpose"): "Table 4.8: Local Application Data",
    ("Layer", "Responsibility", "Example Components"): "Table 5.1: Implementation Layers",
    ("Category", "Technology / Library", "Implementation Purpose"): "Table 5.2: Implementation Technology Stack",
    ("Artifact", "Source Path", "Report Use"): "Table 5.3: Training Notebook Artifacts",
    ("Folder", "Purpose"): "Table 5.4: Runtime Data Folders",
    ("Package", "Main Responsibility"): "Table 5.5: Project Package Responsibilities",
    ("Backend Component", "Main File / Package", "Implementation Role"): "Table 5.6: Backend Components",
    ("Responsibility Group", "Main Tasks"): "Table 5.7: Backend Runtime Responsibilities",
    ("Navigation Group", "Page", "Purpose"): "Table 5.8: Frontend Navigation Groups",
    ("Workflow", "Main Pages"): "Table 5.9: Frontend User Workflows",
    ("Database Area", "Tables / Data", "Purpose"): "Table 5.10: Database Areas",
    ("Table Group", "Tables", "Main Use"): "Table 5.11: Database Table Groups",
    ("Role", "Capabilities"): "Table 5.12: Local Account Roles",
    ("Action Category", "Examples", "Purpose"): "Table 5.13: Alarm Action Categories",
    ("Review Question", "Related Analytics Output"): "Table 5.14: Analytics Review Questions",
    ("Implementation Concern", "Main Module or File", "Report Evidence", "Verification Method"): "Table 5.15: Implementation Traceability Map",
    ("Metric", "Final Value", "Purpose"): "Table 6.1: Final Performance Metrics Plan",
    ("Metric", "Final Value", "Source / Purpose"): "Table 6.1: Final Model Validation Metrics",
    ("Class", "Precision", "Recall", "mAP@0.5", "Notes"): "Table 6.2: Class-Level Evaluation Plan",
    ("Class", "Images", "Instances", "Precision", "Recall", "mAP@0.5", "mAP@0.5:0.95", "Notes"): "Table 6.2: Class-Level Validation Results",
    ("Runtime Measure", "Value", "Notes"): "Table 6.3: Runtime Evaluation Plan",
    ("Metric", "Meaning in Smart Eye", "Risk if Weak"): "Table 6.4: Evaluation Metric Interpretation",
    ("Evaluation Step", "Evidence Produced", "Report Location"): "Table 6.5: Evaluation Data Collection Plan",
    ("Violation Class", "Precision", "Recall", "mAP@0.5", "Interpretation"): "Table 6.6: PPE Violation Class Interpretation",
    ("Dataset Visual", "What to Look For", "Report Meaning"): "Table 6.7: Dataset Figure Interpretation Guide",
    ("Training Visual", "What to Look For", "Report Meaning"): "Table 6.8: Training Figure Interpretation Guide",
    ("Validation Visual", "What to Look For", "Report Meaning"): "Table 6.9: Validation Figure Interpretation Guide",
    ("Qualitative Visual", "What to Compare", "Report Meaning"): "Table 6.10: Qualitative Figure Interpretation Guide",
    ("Application Visual", "Expected Evidence", "Why It Matters"): "Table 6.11: Application Figure Interpretation Guide",
    ("Threat", "Effect on Report Claims", "Mitigation"): "Table 7.1: Threats to Validity",
    ("Error Type", "Example", "Possible Cause"): "Table 7.2: Error Analysis Categories",
    ("Priority", "Improvement", "Reason"): "Table 9.1: Future Work Priorities",
    ("Item", "Required Action", "Expected Result"): "Table D.1: Demonstration Setup Checklist",
    ("Evidence Type", "Screenshot / Output Needed", "Report Section"): "Table D.2: Demonstration Evidence Checklist",
    ("Test ID", "Test Case", "Steps", "Expected Result"): "Table E.1: Functional Test Cases",
    ("Test ID", "Quality Attribute", "Test Method", "Expected Result"): "Table E.2: Non-Functional Test Cases",
    ("Test ID", "Evaluation Focus", "Required Evidence"): "Table E.3: Model Evaluation Test Cases",
    ("Table", "Main Fields", "Purpose"): "Table F.1: Database Table Reference",
    ("Screenshot ID", "Page / Output", "What Must Be Visible"): "Table G.1: Screenshot Capture Checklist",
    ("Screenshot ID", "Required Screenshot", "File Name", "Purpose"): "Table G.1: User Interface Screenshot Checklist",
    ("Result ID", "Required Result Asset", "File Name", "Purpose"): "Table G.2: Result Figure Checklist",
}


def infer_table_caption(table_node: Tag, table_index: int) -> str:
    first_row = table_node.find("tr")
    headers = tuple(inline_text(cell) for cell in first_row.find_all(["th", "td"], recursive=False)) if first_row else ()
    return TABLE_CAPTION_BY_HEADERS.get(headers, f"Table {table_index + 1}: Smart Eye Report Table")


def add_table_caption(document: Document, text: str) -> None:
    text = normalize_text(text)
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    add_tc_field(paragraph, text, "T")


def add_table(document: Document, table_node: Tag, context: dict[str, int]) -> None:
    rows = []
    for tr in table_node.find_all("tr"):
        cells = tr.find_all(["th", "td"], recursive=False)
        if cells:
            rows.append(cells)
    if not rows:
        return
    table_index = context.get("table_index", 0)
    add_table_caption(document, infer_table_caption(table_node, table_index))
    context["table_index"] = table_index + 1
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(columns):
            cell_node = row[col_idx] if col_idx < len(row) else None
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            if cell_node is not None:
                paragraph = cell.paragraphs[0]
                add_inline_runs(paragraph, cell_node, bold=row_idx == 0)
    document.add_paragraph()


def add_list(document: Document, list_node: Tag, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in list_node.find_all("li", recursive=False):
        nested = li.find(["ul", "ol"], recursive=False)
        if nested:
            nested.extract()
        add_paragraph(document, inline_text(li), style=style)


def add_image_placeholder(document: Document, img_node: Tag, base_dir: Path) -> None:
    src = img_node.get("src", "")
    alt = img_node.get("alt", "Image")
    image_path = (base_dir / src).resolve()
    if src and image_path.exists():
        try:
            document.add_picture(str(image_path), width=Inches(6.2))
            add_paragraph(document, alt)
            return
        except Exception:
            pass
    add_paragraph(document, f"[Image placeholder: {alt} - {src}]")


def add_centered_picture(document: Document, image_path: Path, width: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))


def add_caption(document: Document, text: str) -> None:
    text = normalize_text(text)
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    if re.match(r"^Figure\s+(?:[A-Z]\.)?\d", text):
        add_tc_field(paragraph, text, "F")
    elif re.match(r"^Table\s+(?:[A-Z]\.)?\d", text):
        add_tc_field(paragraph, text, "T")


def add_mermaid_to_docx(document: Document, node: Tag, mermaid_dir: Path | None) -> None:
    diagram = node if "mermaid" in node.get("class", []) else node.find(class_="mermaid")
    if diagram is None:
        return
    raw_idx = diagram.get("data-diagram-index", "")
    try:
        idx = int(raw_idx)
    except ValueError:
        idx = -1
    png_path = mermaid_dir / f"mermaid-{idx + 1:02d}.png" if mermaid_dir and idx >= 0 else None
    if png_path and png_path.exists():
        try:
            add_centered_picture(document, png_path, width=6.2)
            caption = DIAGRAM_CAPTIONS[idx] if 0 <= idx < len(DIAGRAM_CAPTIONS) else f"Figure: Smart Eye workflow diagram {idx + 1}."
            add_caption(document, caption)
            return
        except Exception:
            pass
    add_code_block(document, diagram.get_text(), language="mermaid")


def add_math_to_docx(document: Document, node: Tag, math_dir: Path | None) -> None:
    raw_idx = node.get("data-equation-index", "")
    try:
        idx = int(raw_idx)
    except ValueError:
        idx = -1
    png_path = math_dir / f"equation-{idx + 1:02d}.png" if math_dir and idx >= 0 else None
    if png_path and png_path.exists():
        try:
            add_centered_picture(document, png_path, width=4.2)
            return
        except Exception:
            pass
    add_paragraph(document, inline_text(node))


def add_wrapped_code_to_docx(document: Document, node: Tag) -> None:
    pre = node.find("pre")
    code = pre.find("code") if pre else None
    if pre is None:
        return
    label = node.get("data-code-title", "")
    if not label:
        label_node = node.find(class_="code-label")
        label = inline_text(label_node) if label_node else ""
    classes = code.get("class", []) if code else []
    language = ""
    for cls in classes:
        if cls.startswith("language-"):
            language = cls.replace("language-", "", 1)
            break
    add_code_block(document, code.get_text() if code else pre.get_text(), language=language, label=label)


def render_node_to_docx(
    document: Document,
    node: Tag,
    base_dir: Path,
    mermaid_dir: Path | None,
    math_dir: Path | None,
    context: dict[str, int],
) -> None:
    name = node.name.lower()
    if name == "h1":
        add_paragraph(document, inline_text(node), style="Title")
    elif name == "h2":
        add_paragraph(document, inline_text(node), style="Heading 1")
    elif name == "h3":
        add_paragraph(document, inline_text(node), style="Heading 2")
    elif name == "h4":
        add_paragraph(document, inline_text(node), style="Heading 3")
    elif name == "p":
        if node.find("img"):
            caption_text = inline_text(node)
            if caption_text:
                add_rich_paragraph(document, node)
            for img in node.find_all("img"):
                add_image_placeholder(document, img, base_dir)
        else:
            add_rich_paragraph(document, node)
    elif name == "ul":
        add_list(document, node, ordered=False)
    elif name == "ol":
        add_list(document, node, ordered=True)
    elif name == "table":
        add_table(document, node, context)
    elif name == "figure" and node.find(class_="mermaid"):
        add_mermaid_to_docx(document, node, mermaid_dir)
    elif name == "div" and "mermaid" in node.get("class", []):
        add_mermaid_to_docx(document, node, mermaid_dir)
    elif name == "div" and "code-block" in node.get("class", []):
        add_wrapped_code_to_docx(document, node)
    elif name == "div" and "math" in node.get("class", []):
        add_math_to_docx(document, node, math_dir)
    elif name == "pre":
        code = node.find("code")
        classes = code.get("class", []) if code else []
        language = ""
        for cls in classes:
            if cls.startswith("language-"):
                language = cls.replace("language-", "", 1)
        add_code_block(document, code.get_text() if code else node.get_text(), language=language)
    elif name == "blockquote":
        add_rich_paragraph(document, node, style=None)
    elif name == "hr":
        paragraph = document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)


def write_docx(
    markdown_path: Path,
    output_path: Path,
    page_size: str,
    mermaid_dir: Path | None = None,
    math_dir: Path | None = None,
) -> None:
    html_text = markdown_to_html(markdown_path.read_text(encoding="utf-8"))
    soup = BeautifulSoup(html_text, "html.parser")
    body = soup.body or soup
    document = Document()
    configure_docx(document, page_size)
    enable_update_fields_on_open(document)

    auto_sections = {
        "Table of Contents": (
            r'TOC \o "1-3" \h \z \u',
            "Right-click this field in Word and choose Update Field to generate the table of contents.",
        ),
        "List of Figures": (
            r"TOC \f F \h \z",
            "Right-click this field in Word and choose Update Field to generate the list of figures.",
        ),
        "List of Tables": (
            r"TOC \f T \h \z",
            "Right-click this field in Word and choose Update Field to generate the list of tables.",
        ),
    }
    skip_auto_placeholder = False
    context = {"table_index": 0}

    for child in body.children:
        if isinstance(child, NavigableString):
            continue
        if isinstance(child, Tag):
            if skip_auto_placeholder and child.name.lower() == "p":
                text = inline_text(child)
                if text.startswith("To be generated automatically") or text.startswith("Auto-generated"):
                    skip_auto_placeholder = False
                    continue
                skip_auto_placeholder = False

            if child.name.lower() == "h2":
                heading_text = inline_text(child)
                if heading_text in auto_sections:
                    add_paragraph(document, heading_text, style="Heading 1")
                    field_code, placeholder = auto_sections[heading_text]
                    add_auto_field_block(document, field_code, placeholder)
                    skip_auto_placeholder = True
                    continue

            render_node_to_docx(document, child, markdown_path.parent, mermaid_dir, math_dir, context)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def write_pdf_with_weasyprint(html_text: str, output_path: Path, base_dir: Path) -> None:
    from weasyprint import HTML

    output_path.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html_text, base_url=str(base_dir)).write_pdf(str(output_path))


def find_browser_executable() -> str | None:
    candidates = [
        Path.home() / "AppData/Local/Google/Chrome/Application/chrome.exe",
        Path.home() / "AppData/Local/Microsoft/Edge/Application/msedge.exe",
        Path("C:/Program Files/Google/Chrome/Application/chrome.exe"),
        Path("C:/Program Files (x86)/Google/Chrome/Application/chrome.exe"),
        Path("C:/Program Files/Microsoft/Edge/Application/msedge.exe"),
        Path("C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def wait_for_rendered_assets(page, timeout_ms: int = 30000) -> None:
    try:
        page.wait_for_function(
            "() => window.__mermaidReady === true || window.__mermaidReady === 'error' || !document.querySelector('.mermaid')",
            timeout=timeout_ms,
        )
    except Exception:
        pass
    try:
        page.wait_for_function(
            "() => window.__mathReady === true || !document.querySelector('.math.display')",
            timeout=timeout_ms,
        )
    except Exception:
        pass


def render_visual_assets(html_path: Path, mermaid_dir: Path, math_dir: Path) -> tuple[int, int]:
    from playwright.sync_api import sync_playwright

    mermaid_dir.mkdir(parents=True, exist_ok=True)
    math_dir.mkdir(parents=True, exist_ok=True)
    browser_path = find_browser_executable()
    with sync_playwright() as playwright:
        launch_kwargs = {"headless": True}
        if browser_path:
            launch_kwargs["executable_path"] = browser_path
        browser = playwright.chromium.launch(**launch_kwargs)
        page = browser.new_page(viewport={"width": 1400, "height": 1000}, device_scale_factor=2)
        page.goto(html_path.as_uri(), wait_until="networkidle", timeout=60000)
        wait_for_rendered_assets(page)
        diagrams = page.locator("figure.diagram")
        count = diagrams.count()
        rendered_diagrams = 0
        for idx in range(count):
            figure = diagrams.nth(idx)
            box = figure.bounding_box()
            if not box or box["width"] <= 0 or box["height"] <= 0:
                continue
            figure.screenshot(path=str(mermaid_dir / f"mermaid-{idx + 1:02d}.png"))
            rendered_diagrams += 1
        equations = page.locator(".math.display svg")
        equation_count = equations.count()
        if equation_count == 0:
            equations = page.locator(".math.display mjx-container")
            equation_count = equations.count()
        if equation_count == 0:
            equations = page.locator(".math.display")
            equation_count = equations.count()
        rendered_equations = 0
        for idx in range(equation_count):
            equation = equations.nth(idx)
            box = equation.bounding_box()
            if not box or box["width"] <= 0 or box["height"] <= 0:
                continue
            equation.screenshot(path=str(math_dir / f"equation-{idx + 1:02d}.png"))
            rendered_equations += 1
        browser.close()
        return rendered_diagrams, rendered_equations


def write_pdf_with_browser(html_path: Path, output_path: Path, page_size: str) -> None:
    from playwright.sync_api import sync_playwright

    browser_path = find_browser_executable()
    format_name = "Letter" if page_size == "letter" else "A4"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with sync_playwright() as playwright:
        launch_kwargs = {"headless": True}
        if browser_path:
            launch_kwargs["executable_path"] = browser_path
        browser = playwright.chromium.launch(**launch_kwargs)
        page = browser.new_page(viewport={"width": 1280, "height": 1600})
        page.goto(html_path.as_uri(), wait_until="networkidle", timeout=60000)
        wait_for_rendered_assets(page)
        page.emulate_media(media="print")
        page.pdf(
            path=str(output_path),
            format=format_name,
            print_background=True,
            prefer_css_page_size=True,
            margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
        )
        browser.close()


def write_pdf_with_reportlab(markdown_path: Path, output_path: Path, page_size: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        Preformatted,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    html_text = markdown_to_html(markdown_path.read_text(encoding="utf-8"))
    soup = BeautifulSoup(html_text, "html.parser")
    body = soup.body or soup
    pagesize = letter if page_size == "letter" else A4
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=pagesize,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    styles["Normal"].fontName = "Times-Roman"
    styles["Normal"].fontSize = 12
    styles["Normal"].leading = 16
    styles.add(ParagraphStyle("CodeBlock", fontName="Courier", fontSize=8.5, leading=10, backColor=colors.whitesmoke))
    story = []

    for child in body.children:
        if isinstance(child, NavigableString) or not isinstance(child, Tag):
            continue
        name = child.name.lower()
        if name == "h1":
            story.append(Paragraph(html.escape(inline_text(child)), styles["Title"]))
        elif name == "h2":
            story.append(Paragraph(html.escape(inline_text(child)), styles["Heading1"]))
        elif name == "h3":
            story.append(Paragraph(html.escape(inline_text(child)), styles["Heading2"]))
        elif name == "h4":
            story.append(Paragraph(html.escape(inline_text(child)), styles["Heading3"]))
        elif name == "p":
            story.append(Paragraph(html.escape(inline_text(child)), styles["Normal"]))
        elif name in ("ul", "ol"):
            for li in child.find_all("li", recursive=False):
                prefix = "- " if name == "ul" else "1. "
                story.append(Paragraph(html.escape(prefix + inline_text(li)), styles["Normal"]))
        elif name == "table":
            rows = []
            for tr in child.find_all("tr"):
                cells = tr.find_all(["th", "td"], recursive=False)
                if cells:
                    rows.append([Paragraph(html.escape(inline_text(cell)), styles["Normal"]) for cell in cells])
            if rows:
                table = Table(rows, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 5),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ]
                    )
                )
                story.append(table)
        elif name == "pre":
            code = child.find("code")
            story.append(Preformatted(code.get_text() if code else child.get_text(), styles["CodeBlock"]))
        elif name == "div" and "code-block" in child.get("class", []):
            label = child.find(class_="code-label")
            pre = child.find("pre")
            code = pre.find("code") if pre else None
            if label:
                story.append(Paragraph(f"<b><i>{html.escape(inline_text(label))}</i></b>", styles["Normal"]))
            if pre:
                story.append(Preformatted(code.get_text() if code else pre.get_text(), styles["CodeBlock"]))
        elif name == "div" and "math" in child.get("class", []):
            equation_text = inline_text(child).replace("\\[", "").replace("\\]", "")
            story.append(Paragraph(html.escape(equation_text), styles["Normal"]))
        elif name == "figure" and child.find(class_="mermaid"):
            caption = child.find("figcaption")
            if caption:
                story.append(Paragraph(f"<i>{html.escape(inline_text(caption))}</i>", styles["Normal"]))
            story.append(Paragraph("Diagram source fallback. Browser PDF/HTML export renders this as a figure.", styles["Normal"]))
            story.append(Preformatted(child.find(class_="mermaid").get_text(), styles["CodeBlock"]))
        elif name == "hr":
            story.append(PageBreak())
        story.append(Spacer(1, 6))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def pdf_page_count(path: Path) -> int | None:
    try:
        from PyPDF2 import PdfReader

        return len(PdfReader(str(path)).pages)
    except Exception:
        return None


def next_numbered_base_name(out_dir: Path, stem: str) -> str:
    pattern = re.compile(rf"^{re.escape(stem)}_(\d+)$")
    highest = 0
    if out_dir.exists():
        for path in out_dir.iterdir():
            match = pattern.match(path.stem)
            if match:
                highest = max(highest, int(match.group(1)))
    return f"{stem}_{highest + 1}"


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export the Smart Eye Markdown report to DOCX and PDF.")
    parser.add_argument("markdown", nargs="?", default="docs/final_report.md", help="Markdown report path.")
    parser.add_argument("--out-dir", default="docs/report_exports", help="Directory for generated files.")
    parser.add_argument(
        "--base-name",
        default=None,
        help="Output base name. Defaults to the next available numbered name, such as final_report_1.",
    )
    parser.add_argument("--page-size", choices=sorted(PAGE_SIZES), default="letter", help="Page size for DOCX/PDF.")
    parser.add_argument(
        "--pdf-engine",
        choices=["browser", "reportlab", "weasyprint", "auto"],
        default="browser",
        help="PDF renderer. Browser renders CSS and Mermaid best; ReportLab is the fallback.",
    )
    parser.add_argument("--docx-only", action="store_true", help="Export DOCX only and skip PDF generation.")
    parser.add_argument(
        "--no-render-mermaid",
        action="store_true",
        help="Skip rendered diagram and equation screenshots for DOCX. Browser PDF/HTML can still render them.",
    )
    parser.add_argument("--keep-html", action="store_true", help="Keep the generated HTML preview file.")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    markdown_path = Path(args.markdown).resolve()
    if not markdown_path.exists():
        print(f"Markdown file not found: {markdown_path}", file=sys.stderr)
        return 2

    out_dir = Path(args.out_dir).resolve()
    base_name = args.base_name or next_numbered_base_name(out_dir, markdown_path.stem)
    html_path = out_dir / f"{base_name}.html"
    docx_path = out_dir / f"{base_name}.docx"
    pdf_path = out_dir / f"{base_name}.pdf"
    mermaid_dir = out_dir / "assets" / "mermaid"
    math_dir = out_dir / "assets" / "math"

    out_dir.mkdir(parents=True, exist_ok=True)
    html_text = write_html(markdown_path, html_path)
    mermaid_count = 0
    equation_count = 0
    if not args.no_render_mermaid:
        try:
            mermaid_count, equation_count = render_visual_assets(html_path, mermaid_dir, math_dir)
        except Exception as exc:
            print(f"Visual asset rendering failed for DOCX, using source fallback: {exc}", file=sys.stderr)
    write_docx(markdown_path, docx_path, page_size=args.page_size, mermaid_dir=mermaid_dir, math_dir=math_dir)

    pdf_engine = None
    pages = None
    if not args.docx_only:
        if args.pdf_engine == "browser":
            try:
                write_pdf_with_browser(html_path, pdf_path, page_size=args.page_size)
                pdf_engine = "Browser"
            except Exception as exc:
                print(f"Browser PDF export failed, using ReportLab fallback: {exc}", file=sys.stderr)
                write_pdf_with_reportlab(markdown_path, pdf_path, page_size=args.page_size)
                pdf_engine = "ReportLab"
        elif args.pdf_engine == "reportlab":
            write_pdf_with_reportlab(markdown_path, pdf_path, page_size=args.page_size)
            pdf_engine = "ReportLab"
        else:
            try:
                write_pdf_with_weasyprint(html_text, pdf_path, markdown_path.parent)
                pdf_engine = "WeasyPrint"
            except Exception as exc:
                if args.pdf_engine == "weasyprint":
                    raise
                print(f"WeasyPrint PDF export failed, using ReportLab fallback: {exc}", file=sys.stderr)
                write_pdf_with_reportlab(markdown_path, pdf_path, page_size=args.page_size)
                pdf_engine = "ReportLab"
        pages = pdf_page_count(pdf_path)
    if not args.keep_html:
        html_path.unlink(missing_ok=True)

    print(f"DOCX: {docx_path}")
    if not args.docx_only:
        print(f"PDF : {pdf_path}")
        print(f"PDF engine: {pdf_engine}")
    if pages is not None:
        print(f"PDF pages: {pages}")
    if mermaid_count:
        print(f"Mermaid diagrams rendered for DOCX: {mermaid_count}")
    else:
        print("Note: Mermaid diagrams were not rendered into DOCX images; source fallback was used.")
    if equation_count:
        print(f"Equations rendered for DOCX: {equation_count}")
    else:
        print("Note: Equations were not rendered into DOCX images; text fallback was used.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
